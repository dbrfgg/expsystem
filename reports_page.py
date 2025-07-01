from PyQt5.QtWidgets import QWidget, QLabel, QVBoxLayout, QPushButton, QFileDialog
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QMessageBox
from docx import Document
from docx.shared import Inches
import pandas as pd
import os
import tempfile
from docx2pdf import convert
import time


class ReportsPage(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        layout.setAlignment(Qt.AlignTop)

        # Новый современный стиль для верхней части и кнопок
        self.setStyleSheet("""
            QLabel#reportsTitle {
                font-size: 20px;
                font-weight: bold;
                color: #232526;
                margin-bottom: 8px;
            }
            QLabel#reportsDesc {
                font-size: 15px;
                color: #4a4a4a;
                margin-bottom: 18px;
            }
            QPushButton#reportsBtn {
                min-height: 44px;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #f8fafc, stop:1 #e0e7ef);
                color: #232526;
                padding: 12px 22px;
                font-size: 16px;
                border: none;
                border-radius: 10px;
                margin: 8px 0 12px 0;
                text-align: left;
                transition: background 0.3s, color 0.3s, box-shadow 0.3s;
                box-shadow: 0 2px 8px rgba(44,62,80,0.06);
                letter-spacing: 0.5px;
                word-break: break-word;
                white-space: normal;
                line-height: 1.2;
            }
            QPushButton#reportsBtn:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #e0f7fa, stop:1 #b2ebf2);
                color: #00796b;
                box-shadow: 0 4px 16px rgba(26,188,156,0.10);
            }
        """)

        title = QLabel("Вкладка: Отчёты")
        title.setObjectName("reportsTitle")
        layout.addWidget(title)

        self.description = QLabel("Нажмите кнопку ниже, чтобы сгенерировать отчёт в нужном формате.")
        self.description.setObjectName("reportsDesc")
        layout.addWidget(self.description)

        self.word_btn = QPushButton("📝 Сохранить отчёт в Word")
        self.word_btn.setObjectName("reportsBtn")
        self.word_btn.clicked.connect(self.generate_word_report)
        layout.addWidget(self.word_btn)

        self.pdf_btn = QPushButton("📄 Сохранить отчёт в PDF")
        self.pdf_btn.setObjectName("reportsBtn")
        self.pdf_btn.clicked.connect(self.generate_pdf_report)
        layout.addWidget(self.pdf_btn)

        self.setLayout(layout)

    def get_project_data(self):
        """Получает данные проекта из MainWindow"""
        return self.window().get_project_data()

    def get_calculations_data(self):
        """Получает рассчитанные метрики из CalculationsPage"""
        calc_page = self.window().pages.get("Рассчитанные значения")
        if calc_page and hasattr(calc_page, "calculations_data"):
            return calc_page.calculations_data
        return None

    def get_recommendations_data(self):
        """Получает рекомендации из RecommendationsPage"""
        rec_page = self.window().pages.get("Рекомендации")
        if rec_page and hasattr(rec_page, "table"):
            recommendations = []
            for row in range(rec_page.table.rowCount()):
                stage = rec_page.table.item(row, 0).text()
                problem = rec_page.table.item(row, 1).text()
                recommendation = rec_page.table.item(row, 2).text()
                recommendations.append([stage, problem, recommendation])
            return recommendations
        return []

    def get_charts(self):
        """Получает пути к сохранённым графикам из AnalysisPage"""
        analysis_page = self.window().pages.get("Анализ")
        if analysis_page and hasattr(analysis_page, "save_charts"):
            return analysis_page.save_charts()
        return []

    def generate_word_report(self, file_path=None):
        """Генерирует отчёт в формате Word. Если file_path не указан, запрашивает у пользователя."""
        start_time = time.time()
        if not file_path:
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Сохранить отчёт", "", "Word Files (*.docx)"
            )
            if not file_path:
                return False

        try:
            doc = Document()
            doc.add_heading("Отчёт по анализу проекта", 0)

            # Раздел 1: Данные проекта
            t1 = time.time()
            project_data = self.get_project_data()
            if project_data is not None:
                doc.add_heading("1. Данные проекта", level=1)
                table = doc.add_table(rows=1 + len(project_data), cols=len(project_data.columns))
                table.style = "Table Grid"
                for j, col in enumerate(project_data.columns):
                    table.rows[0].cells[j].text = col
                for i, row in project_data.iterrows():
                    for j, val in enumerate(row):
                        table.rows[i + 1].cells[j].text = str(val)
            else:
                doc.add_paragraph("Данные проекта отсутствуют.")
            print(f"Project data section took {time.time() - t1:.2f} seconds")

            # Раздел 2: Рассчитанные метрики
            t2 = time.time()
            calc_data = self.get_calculations_data()
            if calc_data is not None:
                doc.add_heading("2. Рассчитанные метрики", level=1)
                selected_columns = ["Этап", "ΔT", "ΔC", "E"]
                filtered_calc_data = calc_data[selected_columns]
                table = doc.add_table(rows=1 + len(filtered_calc_data), cols=len(filtered_calc_data.columns))
                table.style = "Table Grid"
                for j, col in enumerate(filtered_calc_data.columns):
                    table.rows[0].cells[j].text = col
                for i, row in filtered_calc_data.iterrows():
                    for j, val in enumerate(row):
                        table.rows[i + 1].cells[j].text = str(val)
            else:
                doc.add_paragraph("Рассчитанные метрики отсутствуют.")
            print(f"Calculations section took {time.time() - t2:.2f} seconds")

            # Раздел 3: Рекомендации
            t3 = time.time()
            recommendations = self.get_recommendations_data()
            if recommendations:
                doc.add_heading("3. Рекомендации", level=1)
                table = doc.add_table(rows=1 + len(recommendations), cols=3)
                table.style = "Table Grid"
                table.rows[0].cells[0].text = "Этап"
                table.rows[0].cells[1].text = "Проблема"
                table.rows[0].cells[2].text = "Рекомендация"
                for i, (stage, problem, recommendation) in enumerate(recommendations):
                    table.rows[i + 1].cells[0].text = stage
                    table.rows[i + 1].cells[1].text = problem
                    table.rows[i + 1].cells[2].text = recommendation
            else:
                doc.add_paragraph("Рекомендации отсутствуют.")
            print(f"Recommendations section took {time.time() - t3:.2f} seconds")

            # Раздел 4: Графики
            t4 = time.time()
            charts = self.get_charts()
            if charts:
                doc.add_heading("4. Графики", level=1)
                for chart_name, chart_path in charts:
                    doc.add_paragraph(chart_name)
                    doc.add_picture(chart_path, width=Inches(5.5))
            else:
                doc.add_paragraph("Графики отсутствуют.")
            print(f"Charts section took {time.time() - t4:.2f} seconds")

            doc.save(file_path)
            # Очистка временных файлов
            t5 = time.time()
            for _, chart_path in charts:
                os.remove(chart_path)
            print(f"File cleanup took {time.time() - t5:.2f} seconds")

            # Показываем сообщение об успешном сохранении
            QMessageBox.information(self, "Успех", "Файл Word успешно сохранён!")
            print(f"Total Word report generation took {time.time() - start_time:.2f} seconds")
            return True
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось создать отчёт: {str(e)}")
            return False

    def generate_pdf_report(self):
        """Генерирует отчёт в формате PDF путём преобразования Word-документа"""
        start_time = time.time()
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить отчёт", "", "PDF Files (*.pdf)"
        )
        if not file_path:
            return

        try:
            # Создаём временный Word-файл
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
                temp_docx_path = temp_docx.name

            # Генерируем Word-документ во временный файл
            t1 = time.time()
            if not self.generate_word_report(temp_docx_path):
                return
            print(f"Word file generation for PDF took {time.time() - t1:.2f} seconds")

            # Преобразуем Word в PDF
            t2 = time.time()
            convert(temp_docx_path, file_path)
            print(f"PDF conversion took {time.time() - t2:.2f} seconds")

            # Удаляем временный Word-файл
            t3 = time.time()
            os.remove(temp_docx_path)
            print(f"Temp file cleanup took {time.time() - t3:.2f} seconds")

            QMessageBox.information(self, "Успех", "Файл PDF успешно сохранён!")
            print(f"Total PDF report generation took {time.time() - start_time:.2f} seconds")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось создать PDF: {str(e)}")
